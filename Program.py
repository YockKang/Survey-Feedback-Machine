##################
## Import Files ##
##################

import os
import sys
import time
import nltk
import sumy
import numpy
from sumy.parsers.html import HtmlParser
from sumy.parsers.plaintext import PlaintextParser
from sumy.nlp.tokenizers import Tokenizer
from sumy.nlp.stemmers import Stemmer
from sumy.utils import get_stop_words
from sumy.summarizers.lex_rank import LexRankSummarizer
import numpy as np
import matplotlib.pyplot as plt
from matplotlib.ticker import MaxNLocator
from collections import namedtuple
import csv
import docx
from docx.text.paragraph import Paragraph




#############################
## NUMERICAL/DATA ANALYSIS ##
#############################

def count_num(lst,num):
    '''function that counts number of occurences for a particular object in a list'''
    count=0
    for i in range(len(lst)):
        if lst[i]==num:
            '''Checks if it is the number specified and adds'''
            count+=1
    return count

    
def percentage(lst,num, sf=2):
    '''gives perecentage result'''
    result=round(count_num(lst,num)/len(lst),sf)
    return result

def remove_duplicates(lst):
    '''ensures not empty'''
    if lst!=[]:
        alst=sorted(lst)
        outlst=[]
        outlst.append(alst[0])
        '''above sorts the items, below checks if before is different, if Y, add to outlst'''
        for i in range(1,len(alst)):
            if alst[i-1]!=alst[i]:
                outlst.append(alst[i])
        return outlst
    else:
        return []

      
def highest_num_index(nums):
    indexes=[]
    highest=max(nums)
    '''above finds highest numbers, below fids the index of these numbers'''
    for i in range(len(nums)):
        if nums[i]==highest:
            indexes.append(i)
    return indexes

  
def mean(nums,round_to=2):
    '''find mean, round to 2'''
    if nums==[]:
        return 0
    total=0
    for i in nums:
        total+=i
    if round_to==0:
        return int(total/len(nums))    
    return round(total/len(nums),round_to)

def median(nums):
    if nums==[]:
        return 'NA'
    sortednums=sorted(nums)
    '''above sorts numbers, below finds middle to get median'''
    if len(nums)%2==1:
        return sortednums[int(len(sortednums)/2-0.5)]
    else:
        return (sortednums[int(len(sortednums)/2)]+
            sortednums[int(len(sortednums)/2-1)])/2

      
def mode(nums):
    if nums==[]:
        return 'NA'
    no_dupe=remove_duplicates(nums)
    count=[]
    '''above finds non-duplicated contents, below finds corresponding indexes and returns mode'''
    for i in range(len(no_dupe)):
        count.append(count_num(nums,no_dupe[i]))
    mode_indexes=highest_num_index(count)
    mode=[]
    for i in range(len(mode_indexes)):
        mode.append(no_dupe[mode_indexes[i]])

    
    return mode[0]

  
def pop_std_dev(nums,round_to=3):
    if nums==[]:
        return 'NA'
    '''Step 1: find mean'''
    avg=0
    avg=mean(nums,20)
    '''Step 2: find square of difference for all'''
    sq_diff=[]
    for i in range(len(nums)):
        sq_diff.append((nums[i]-avg)**2)
    '''Step 3: find mean of sq_diff'''
    pop_variance=mean(sq_diff,20)
    '''Step 4: sqrt variance and done!'''
    std_dev=pop_variance**0.5
    return round(std_dev,round_to)

  
def smpl_std_dev(nums,round_to=3):
    if nums==[]:
        return 'NA'
    '''Step 1: find mean'''
    avg=0
    avg=mean(nums,20)
    '''Step 2: find square of difference for all'''#Step 2: find square of difference for all
    sq_diff=[]
    for i in range(len(nums)):
        sq_diff.append((nums[i]-avg)**2)
    '''Step 3: find sum of sq_diff divided by (num of population)-1'''
    total=0
    for i in sq_diff:
        total+=i
    smpl_variance=total/(len(sq_diff)-1)
    '''Step 4: sqrt variance and done!'''
    std_dev=smpl_variance**0.5
    return round(std_dev,round_to)



def correlation(datax,datay,qx,qy):
    '''use pearson correlation to determine relationship between questions'''
    datax=strtonums(datax)
    datay=strtonums(datay)
    if len(datay)!=len(datax):
        errormsg='The number of responses on\nColumn '+chr(65+qx)+' and\nColumn '+chr(65+qy)+' \nAre not equal\n(i.e. there is a non-numerical response in one of the questions)'
        error(errormsg)
    meanx=mean(datax)
    meany=mean(datay)

    sample_stdx=smpl_std_dev(datax)
    sample_stdy=smpl_std_dev(datay)

    r=0#s(coefficient)
    temp_counter=0
    for i in range(len(datax)):
        temp_counter= ((datax[i]-meanx)/sample_stdx) * ((datay[i]-meany)/sample_stdy)
        r+=temp_counter
        temp_counter=0
    return r/(len(datax)-1)



##########
## SUMY ##
##########

def summarize(str):
    '''uses LexRankSummariser to summarise the open-ended questions'''
    returnstring=""
    file=open('condensethis.txt','w+')
    file.writelines(str)
    file.close()
    parser = PlaintextParser.from_file('condensethis.txt', Tokenizer("english"))
    
    summarizer = LexRankSummarizer()
    summary =summarizer(parser.document,1)
    for sentence in summary:
        returnstring+="\n"+sentence._text
    return returnstring




#####################################
##  MATPLOTLIB DATA VISUALISATION  ##
#####################################

def bar(num, question, option_no, options, xaxis, yaxis, data, std_men = (2, 3, 4, 1, 2),
        bar_width = 0.35,opacity = 0.4):
    '''creates a bar graph'''
    fig, ax = plt.subplots()

    index = np.arange(option_no)

    filename = 'Q' + str(num) + '.png'
    
    error_config = {'ecolor': '0.3'}

    rects1 = ax.bar(index, data, bar_width,
                    alpha=opacity, color='r', 
                    error_kw=error_config, label="Options")

    ax.set_xlabel(xaxis)
    ax.set_ylabel(yaxis)
    ax.set_title(question)
    ax.set_xticks(index + bar_width / 2)
    ax.set_xticklabels(options)
    ax.legend()

    fig.tight_layout()
    plt.savefig(filename)
    
    
def pie(numbers,options,num,question):
    '''creates a pie chart'''
    filename='Q'+str(num)+'.png'
    fig1, ax1 = plt.subplots()
    ax1.pie(numbers, labels=options, autopct='%1.1f%%',
            shadow=True, startangle=90)
    ax1.axis('equal')  # Equal aspect ratio ensures that pie is drawn as a circle.
    ax1.set_title(question)

    plt.savefig(filename)


def scatterplot(x,y,question1,question2):
    '''creates a scatterplot graph with linear trendline'''
    filename='Scatterplot'
    x=strtonums(x)
    y=strtonums(y)
    plt.scatter(x,y)
    # calc the trendline (best fit line)
    z = numpy.polyfit(x, y, 1)
    p = numpy.poly1d(z)
    plt.plot(x,p(x))
    plt.savefig(filename)
    
    
##################
## FILE READING ##
##################

def read_categories(filename):
    '''reads headers'''
    file=open(filename,"r")
    categories=file.readlines(1)
    
    for i in categories:
        i.strip()
        categories=i.split(",")
    #print(categories)
    return categories
    file.close()

def read_data(filename):
    '''uses csv to find the separate variables'''
    file=csv.reader(open(filename, 'r'), delimiter=',')
    alldata=[]
    for line in file:
        alldata.append(line)
    categories=read_categories(filename)
    
    data=[]
    for i in range(len(categories)):
        temp=[]
        for j in range(1,len(alldata)): #omit the first line of questions
            line=alldata[j]
            try:
                temp.append(line[i])
            except:
                temp.append("")
        data.append(temp)
    return data

def read_config(filename):
    '''reads Config.csv to check for key question types'''
    file=csv.reader(open(filename, 'r'), delimiter=',')
    alldata=[]
    for line in file:
        alldata.append(line)
    data=[]
    temp=[]
    
    for i in alldata:
        if i[0]=='Admin' or i[0]=='Qualitative':
            temp2=[]
            temp2.append(i[0])
            data.append(temp2)
        elif i[0]== 'Num':
            
            try:
                firstnum=int(i[1])
                lastnum=int(i[2])
            except ValueError:
                errormsg='We have encountered an error in Config.csv, row: ' + str(alldata.index(i)+1)+'\nA Num question cointains a field that is not a number'
                error(errormsg)
            for num in range(firstnum,lastnum+1):
                temp.append(num)
            data.append(['Num',temp])
            temp=[]
        elif i[0]=='Categorical':
            temp.append(i[0])
            options=[]
            for option_num in range(1,len(i)):
                if i[option_num]!="":
                    options.append(i[option_num])
            temp.append(options)
            data.append(temp)
            temp=[]
        elif i[0]=='':
            filler=0
        else:
            errormsg='The question class (Column 1) of the question on row '+str(alldata.index(i)+1)+' is not under any of the 4 classes:\nNum\nCategrical\nAdmin\nQualitative'
            error(errormsg)
    return data

def write(filename,data):
    '''writes data into filename'''
    f=open(filename,"w+")
    f.writelines(data)
    f.close()


def read_graph(filename):
    '''reads graph config for pie chart or bar graph'''
    file=csv.reader(open(filename, 'r'), delimiter=',')
    alldata=[]
    for line in file:
        alldata.append(line)
    data=[]
    temp=[]
    for i in range(len(alldata)):
        if alldata[i-1][1]=="Graph Type(Either Pie or Bar)":
            if alldata[i][1]=="Bar" or "Pie":
                return alldata[i][1]
            else:
                return "Bar"
    return 'Bar'


def read_executive():
    '''reads options for executive function for Correlation or Difference'''
    file=csv.reader(open("Config.csv", 'r'), delimiter=',')
    alldata=[]
    for line in file:
        alldata.append(line)
    
    data=[]
    
    
    for i in range(len(alldata)):
        if alldata[i-1][1]=="Executive Analysis":
            if alldata[i][1]=="Correlation":
                
                data.append("Correlation")
                data.append(int(alldata[i][2])-1)
                data.append(int(alldata[i][3])-1)
            elif alldata[i][1]=="Difference":
                
                data.append("Difference")
                data.append(int(alldata[i][2])-1)
                data.append(int(alldata[i][3])-1)
    
    return data

def read_summary():
    '''reads option for summary of program'''
    file=csv.reader(open("Config.csv", 'r'), delimiter=',')
    alldata=[]
    for line in file:
        alldata.append(line)
    data=[]
    
    
    for i in range(len(alldata)):
        if alldata[i-1][1]=="Summary for perception of program?(Y or N)":
            if alldata[i][1]=="Y":
                return 1
            else:
                return 0


def read_demand():
    '''reads for the reqiremants to add to the demand'''
    file=csv.reader(open("Config.csv", 'r'), delimiter=',')
    alldata=[]
    for line in file:
        alldata.append(line)
    out=[]
    options=[]
    for i in range(len(alldata)):
        if alldata[i-1][1]=="Demand Question Number":
            if alldata[i][1]=='':
                out=['',[]]
                break
            question=alldata[i][1]
            out.append(question)
            for j in range(2,len(alldata[i])):
                if alldata[i][j]!='':
                    options.append(alldata[i][j])
            out.append(options)
    return out


        
##########################
## Categorical Analysis ##
##########################

def analyse_summary(data,config):
    """This algorith compares the mean and the standard deviation. It is slightly more accurate than using mode"""
    if read_summary()==1:
        rangeofchoices=[]


        perception=""

        for i in range(len(config)):
            if config[i][0]=="Num":
                choiceslist=[strtonums(config[i][1])[0],
                             strtonums(config[i][1])[len(config[i][1])-1]]
                centerpoint=(choiceslist[1]-choiceslist[0])/2
                #Where equal number of options lie on both sides of this number
                #eg for 1,2,3,4,5 centerpoint is 3, where there are 2 options on both sides of 3
                meaan=mean(strtonums(data[i]))
                std_dev=pop_std_dev(strtonums(data[i]))
               
                if meaan>centerpoint:
                    if meaan-std_dev>centerpoint-std_dev/2:
                        rangeofchoices.append("Positive")
                    elif meaan-std_dev<=centerpoint-std_dev/2:
                        rangeofchoices.append("Positive to Mixed")
                elif meaan==centerpoint:
                    if meaan+std_dev<centerpoint+std_dev/2:
                        rangeofchoices.append("Negative")
                    elif meaan+std_dev>=centerpoint+std_dev/2:
                        rangeofchoices.append("Negative to Mixed")


        
        pos=count_num(rangeofchoices,"Positive")
        postomix=count_num(rangeofchoices,"Positive to Mixed")
        negtomix=count_num(rangeofchoices,"Negative to Mixed")
        neg=count_num(rangeofchoices,"Negative")

        if pos+postomix>negtomix+neg:
            if pos>postomix:
                perception="Positive"
            else:
                perception="Positive to Mixed"
        elif pos+postomix==negtomix+neg:
            perception="Mixed"
        else: #if negtomix+neg=pos+postomix
            if neg>negtomixed:
                perception="Negative"
            else:
                perception="Negative to Mixed"
        

        write("Output.txt",
              "Summary TL:DR (Too long: Didn't read)\n\nThere is an overall "+perception+" response in the Numerical Questions in the survey.\n")
        add_text()                        
            
                    


        
def analyse_numerical(data, index,options):
    """Analyses NUMERICAL questions and returns string(to add to "Output.docx")"""
    returnstring="\nQuestion "+str(index)+":\n"
    x=data[index-1]
    x=strtonums(x)
    returnstring+="Total Responses : "+str(len(x))+"\n"
    for option in options:
        returnstring+="Responses for "+str(option)+" : "+str(count_num(x,option))+'\n'
    
    returnstring+="Mean: " +str(mean(x))+ "\nMode: " + str(mode(x))
    returnstring+="\nMedian: " + str(median(x))+"\nStandard Deviation: "+ str(pop_std_dev(x))+"\n"
    return returnstring
            

def strtonums(lst):
    """Converts a list of string ints to a list of ints
    useful as csv reads ints as strings"""
    
    lst2=[]
    for i in lst:
        try:
            lst2.append(int(i))
        except:
            pass
    return lst2
  
def analyse_categorical(data, index, categories):
    """Analyses Categorical questions and returns string(to add to "Output.docx")"""
    returnstring=""
    
    qndata=data[index]
    responsenumbers=[]
    percentages=[]
    for category in categories:
        responsenumbers.append([category,count_num(qndata,category)])
        percentages.append([category,percentage(qndata,category)])
    returnstring+="\nQuestion "+ str(index+1)+ " :\n"
    returnstring+="Total Responses : "+ str(len(qndata))+"\n\n"
    for i in range (len(percentages)):
        returnstring+="Number of responses for '"+str(responsenumbers[i][0])+"' : "+str(responsenumbers[i][1])+"\nPercentage : "+str(percentages[i][1]*100)+"%\n"
    return returnstring

def analyse_qualitative(data, index):
    """Analyses Qualitative questions and returns string(to add to "Output.docx")"""
    returnstring=""
    returnstring+="\nQuestion " + str(index+1) + " Summary :"
    qndata=data[index]
    linedata=''
    for i in range(len(qndata)):
        linedata+=qndata[i]+'\n'
    returnstring+=summarize(linedata)
    returnstring+='\n'
    return returnstring

def analyse_choicenonum(data, index,options):
    '''Function used in visualisation function
    counts responses for each option in numerical questions'''
    x=data[index]
    x=strtonums(x)
    lst=[]
    for option in options:
        lst.append(count_num(x,option))
    
    return lst  

def analyse_choicenocat(data, index, categories): 
    '''Function used in the visualisation function
    Counts the responses for each option for each category for categorical questions'''
    qndata=data[index]
    responsenumbers=[]
    percentages=[]
    for category in categories:
        responsenumbers.append([category,count_num(qndata,category)])
        percentages.append([category,percentage(qndata,category)])

    lst=[]
    for i in range (len(percentages)):
        lst.append(responsenumbers[i][1])
    return lst




def analyse_executive():
    """Central Function that reads the category of Executive Summary, analyses and writes to docx"""
    config=read_executive()
    data=read_data("Responses.csv")
    questions=read_categories("Responses.csv")
    try:
        if config[0]=="Correlation":
            analyse_correlation(data,config,questions)
        elif config[0]=="Difference":
            analyse_difference(data,config,questions)
    except:
        pass
        
        


#minor functions
def relationship(h_or_l,qn1,qn2):
    returnstring= "\n\nThe higher the response for '" +qn1+ "', the " + h_or_l +" the response for '" +qn2+ "'."
    return returnstring

def strongorweak(strengthandpositivity):
    returnstring= "\nThere is a " + strengthandpositivity +" correlation between the results of the two questions.\n"
    return returnstring
#improves readabillity below


def analyse_correlation(data,config,questions):
    '''uses pearson correlation to give a generalised statement of how well the linear trendline plotted correlates to the data'''
    qnno1=config[1]
    qnno2=config[2]
    corr_cof=correlation(data[qnno1],data[qnno2],qnno1,qnno2)
    scatterplot(data[qnno1],data[qnno2],questions[qnno1+1],questions[qnno2+1])

    qn1=questions[qnno1]
    qn2=questions[qnno2]
    returnstring="\nExecutive Summary"
    if corr_cof<-0.5: #Strong negative correlation
        returnstring+= relationship("lower", qn1,qn2)+strongorweak("strong negative")
    elif corr_cof<0: #Weak negative correlation
        returnstring+= relationship("lower", qn1,qn2)+strongorweak("weak negative")
    elif corr_cof==0: #No relationship
        returnstring+= "There is no relationship between the results of these two questions"
    elif corr_cof<0.5: #Weak Positive correlation
        returnstring+= relationship("higher", qn1,qn2)+strongorweak("weak positive")
    elif corr_cof<=1: #Strong Positive
        returnstring+= relationship("higher", qn1,qn2)+strongorweak("strong posiitve")
    write("Output.txt",returnstring)
    add_text()
    add_pic("Scatterplot.png")
        
def analyse_demand(req,data):
    '''finds if in the requirements ans in the responses of the given question and adds one per response'''
    if req[0]!='':
        options=req[1]
        demand=0
        for j in data[int(req[0])-1]:
            for k in options:
                if j==k:
                    demand+=1
        text='\nYour expected demand will be: '+str(demand)+' of '+str(len(data[0]))
        doc = docx.Document('Output.docx')
        doc.add_paragraph(text)
        doc.save('Output.docx')
                        
    
    

def analyse_difference(data,config,questions):
    '''Central Function for the "difference" function under executive summary'''
    data1=strtonums(data[config[1]])
    data2=strtonums(data[config[2]])
    lst=[]
    if len(data1)!=len(data2):
        errormsg='The number of responses for the 2 questions are not equal\n(i.e. there is a non-numerical response in one of the questions)'
        error(errormsg) 
    #code to check length of data
    
    for i in range(len(data1)):
        difference=data1[i]-data2[i]
        if difference>1:
            lst.append("higher")
        elif difference>=-1:
            lst.append("similar")
        elif difference<-1:
            lst.append("lower")
      
    high=count_num(lst,"higher")
    sim=count_num(lst,"similar")
    lower=count_num(lst,"similar")
    i+1

    if high>lower:
        if high>sim:
            writestring="mostly higher"
        elif high<=sim:
            writestring="mostly similar or higher"

    elif lower==high:
        writestring=="mostly similar"
    elif lower>high:
        if lower>sim:
            writestring="mostly lower"
        elif lower<=sim:
            writestring="mostly lower or similar"
    
    write("Output.txt",
          "\nExecutive summary\n\nRespondents generally answered " +writestring+ " for '" +
          questions[config[1]]+"' than '" + questions[config[2]]+"' .")
    add_text()
    
          
        
    
            
            
            
    
################################
## Convert to .docx (MS Word) ##
################################

def delete_paragraph(paragraph):
    '''deletes a paragraph'''
    p = paragraph._element
    p.getparent().remove(p)
    p._p = p._element = None
    
def clear():
    '''clear Output.docx'''
    file=open('Output.txt')
    lines=file.readlines()
    file.close
    doc = docx.Document('Output.docx')
    for i in range(len(doc.paragraphs)):
        delete_paragraph(doc.paragraphs[0])
    doc.save('Output.docx')

def add_text():
    '''adds text fron Output.txt'''
    try:                    
        file=open('Output.txt')
        lines=file.readlines()
        file.close
        doc = docx.Document('Output.docx')
        for j in range(len(lines)):
            line=lines[j].rstrip('\n')
            doc.add_paragraph(line)
        doc.save('Output.docx')
    except PermissionError:
        print("Please close 'Output.docx' before running the program")
        time.sleep(15)
        sys.exit()
        
def add_pic(pic_filename):
    ''' adds picture'''
    doc=docx.Document('Output.docx')
    doc.add_picture(pic_filename,width=docx.shared.Inches(6),height=docx.shared.Inches(4.5))
    doc.save('Output.docx')




##################
## Data Testing ##
##################

def analyse(datafile,config, output):
    '''central function that analyses all the individual questions'''
    returnstring=""
    #code to read file
    questions=read_categories(datafile)
    data=read_data(datafile)
    
    #code to read config folder
    
    #recursive code that directs to analysis functions
    for i in range(len(config)):
        returnstring=""
        cat=config[i][0]
        if cat=="Admin":
            returnstring+="\nQuestion " + str(i+1)+" is a "+cat+" question."+'\n'
        elif cat=="Qualitative":
            
            returnstring+=analyse_qualitative(data,i)
        elif cat=="Num":
            returnstring+=analyse_numerical(data,i+1,config[i][1])
        elif cat=="Categorical":
            returnstring+=analyse_categorical(data,i,config[i][1])
            
        write("Output.txt", returnstring)
        add_text()
        try:
            add_pic('Q'+str(i+1)+'.png')
        except:
            pass #filler code

def visualisation(datafile, qnconfig, graphconfig="Bar"):
    '''central function that plots all the graphs for numerical and categorical question'''
    questionnames=read_categories(datafile)
    data=read_data(datafile)
    
    if graphconfig== "Bar" :
        for i in range(len(qnconfig)):
            if qnconfig[i][0]=="Categorical":
                
                bar(i+1,
                    "Q"+str(i+1)+" : "+questionnames[i],
                    len(qnconfig[i][1]),
                    qnconfig[i][1], 'Options',
                    "Number of Choices",
                    analyse_choicenocat(data,i,qnconfig[i][1]))
                
            elif qnconfig[i][0]=="Num":
                bar(i+1,
                    "Q"+str(i+1)+" : "+questionnames[i],
                    len(qnconfig[i][1]),
                    qnconfig[i][1], 'Options',
                    "Number of Choices",
                    analyse_choicenonum(data,i,qnconfig[i][1]))

    elif graphconfig=="Pie":
        for i in range(len(qnconfig)):
                if qnconfig[i][0]=="Categorical":
                    
                    pie(analyse_choicenocat(data,i,qnconfig[i][1]),
                        qnconfig[i][1],i+1,"Q"+str(i+1)+" : "+questionnames[i])
                    
                elif qnconfig[i][0]=="Num":
                    pie(analyse_choicenonum(data,i,qnconfig[i][1]),
                        qnconfig[i][1],i+1,"Q"+str(i+1)+" : "+questionnames[i])
    

def deleteallpictures(qnconfig):
    '''removes pictures from file directory'''
    for i in range(len(qnconfig)):
        try:
            os.remove("Q"+str(i+1)+".png")
        except:
            pass
    try:
        os.remove("Scatterplot.png")
    except:
        pass

                    

####################
## Error Handling ##
####################

def clear_errors():
    '''wipes Error.txt clean'''
    file=open('Error.txt','w+')
    file.writelines('')
    file.close()
    
def error(message):
    '''function to give error'''
    file=open('Error.txt','w+')
    text=file.readlines()
    print('There is an error, view Error.txt to identify error')
    file.writelines(message)
    file.close()
    time.sleep(15)
    sys.exit()




################
## Test Cases ##
################
        
def start():
    '''runs operations'''
    clear_errors()
    clear()
    Question_config=read_config("Config.csv")
    data=read_data("Responses.csv")
    analyse_summary(data,Question_config)
    analyse_demand(read_demand(),data)
    analyse_executive()
    visualisation("responses.csv", Question_config,read_graph("Config.csv"))
    analyse("responses.csv",Question_config,"Output.txt")
    write("Output.txt","") #clears file
    deleteallpictures(Question_config)
    

start()
